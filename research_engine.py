import os
import shutil
import logging
from typing import List, Dict, Any, Optional
import google.generativeai as genai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
import PyPDF2
import docx
from bs4 import BeautifulSoup
import requests
from dotenv import load_dotenv

load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class IngestionService:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        text = ""
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
        return text

    @staticmethod
    def extract_text_from_url(url: str) -> str:
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(url, timeout=10, headers=headers)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            # Remove scripts and styles
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n')
            return text
        except Exception as e:
            logger.error(f"Error extracting URL: {e}")
            raise ValueError(f"Failed to process URL: {str(e)}")

class ResearchEngine:
    def __init__(self, api_key: str, persist_directory: str = "./data/chroma_db"):
        self.api_key = api_key
        self.persist_directory = persist_directory
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-2.5-flash')
        
        # Load embedding model (cached)
        self.embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # Initialize Vector Store
        os.makedirs(persist_directory, exist_ok=True)
        self.vector_db = Chroma(
            persist_directory=persist_directory,
            embedding_function=self.embeddings,
            collection_name="research_docs"
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100
        )

    def ingest_file(self, file_path: str, filename: str) -> str:
        ext = filename.split(".")[-1].lower()
        if ext == "pdf":
            text = IngestionService.extract_text_from_pdf(file_path)
        elif ext == "docx":
            text = IngestionService.extract_text_from_docx(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        
        chunks = self.text_splitter.split_text(text)
        documents = [
            Document(page_content=chunk, metadata={"source": filename})
            for chunk in chunks
        ]
        
        self.vector_db.add_documents(documents)
        self.vector_db.persist()
        return f"Successfully ingested {len(documents)} chunks from {filename}"

    def ingest_url(self, url: str) -> str:
        text = IngestionService.extract_text_from_url(url)
        chunks = self.text_splitter.split_text(text)
        documents = [
            Document(page_content=chunk, metadata={"source": url})
            for chunk in chunks
        ]
        
        self.vector_db.add_documents(documents)
        self.vector_db.persist()
        return f"Successfully ingested {len(documents)} chunks from {url}"

    def query(self, user_query: str, limit: int = 5) -> Dict[str, Any]:
        try:
            # 1. Retrieve context
            docs = self.vector_db.similarity_search(user_query, k=limit)
            context = "\n\n".join([f"Source: {d.metadata['source']}\nContent: {d.page_content}" for d in docs])
            
            # 2. Generate response via Gemini
            prompt = f"""
            You are an expert Research Analyst. Use the following context to answer the user query.
            Always cite your sources using [Source Name].
            If the context doesn't contain the answer, say "I don't have enough information in my database."
            
            Context:
            {context}
            
            Query: {user_query}
            
            Answer:
            """
            
            response = self.model.generate_content(prompt)
            
            return {
                "answer": response.text,
                "sources": [d.metadata['source'] for d in docs],
                "snippets": [d.page_content for d in docs]
            }
        except Exception as e:
            logger.error(f"Error in query: {e}")
            raise Exception(f"Research engine failed: {str(e)}")

    def get_stats(self) -> Dict[str, Any]:
        try:
            count = self.vector_db._collection.count()
            return {"document_chunks": count}
        except:
            return {"document_chunks": 0}

    def clear_database(self):
        try:
            # Delete the collection properly
            self.vector_db.delete_collection()
            # Re-initialize
            self.vector_db = Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings,
                collection_name="research_docs"
            )
            # Ensure folder is clean if possible, but delete_collection is usually enough
            if os.path.exists(self.persist_directory):
                shutil.rmtree(self.persist_directory)
                os.makedirs(self.persist_directory, exist_ok=True)
            logger.info("Database cleared successfully.")
        except Exception as e:
            logger.error(f"Error clearing DB: {e}")
            raise Exception(f"Failed to clear database: {str(e)}")
